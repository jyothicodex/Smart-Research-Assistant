import os
import io
import time
import streamlit as st
from docx import Document
from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.styles import getSampleStyleSheet
import PyPDF2

# NOTE: if you're using OpenAI's new Python package, adjust imports.
# This code expects: from openai import OpenAI and client.chat.completions.create
try:
    from openai import OpenAI
except Exception:
    OpenAI = None

# -------------------------
# Configuration
# -------------------------
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
MODEL_NAME = "gpt-4o-mini"  # change if needed / available
INITIAL_CREDITS = 100.0
COST_PER_REPORT = 1.0  # mock credit cost

# -------------------------
# Initialize OpenAI client (or run mock)
# -------------------------
client = None
if OPENAI_API_KEY and OpenAI is not None:
    try:
        client = OpenAI(api_key=OPENAI_API_KEY)
    except Exception:
        client = None

# -------------------------
# Streamlit page config
# -------------------------
st.set_page_config(
    page_title="📚 Smart Research Assistant",
    layout="wide",
    initial_sidebar_state="expanded"
)

# -------------------------
# Session state initialization
# -------------------------
if "questions" not in st.session_state:
    st.session_state.questions = 0
if "reports" not in st.session_state:
    st.session_state.reports = 0
if "credits_used" not in st.session_state:
    st.session_state.credits_used = 0.0
if "credits_remaining" not in st.session_state:
    st.session_state.credits_remaining = float(INITIAL_CREDITS)
if "billing_log" not in st.session_state:
    st.session_state.billing_log = []  # list of dicts: {q, cost, ts}
if "live_feed" not in st.session_state:
    # mock live feed: each entry is dict {id, title, source, content, ts}
    st.session_state.live_feed = []
if "sources" not in st.session_state:
    st.session_state.sources = []  # list of strings
if "last_report" not in st.session_state:
    st.session_state.last_report = None

# -------------------------
# Helpers: file extraction & downloads
# -------------------------
def save_to_docx(report_text, title="research_report"):
    doc = Document()
    # add title
    doc.add_heading(title, level=1)
    for line in report_text.split("\n"):
        doc.add_paragraph(line)
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def save_to_pdf(report_text, title="research_report"):
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer)
    styles = getSampleStyleSheet()
    story = [Paragraph(f"<b>{title}</b>", styles["Heading1"])]
    # split into paragraphs
    for paragraph in report_text.split("\n\n"):
        story.append(Paragraph(paragraph.replace("\n", "<br/>"), styles["Normal"]))
    doc.build(story)
    buffer.seek(0)
    return buffer

def extract_file_content(uploaded_files):
    """Extract text from uploaded file-like Streamlit files (pdf/docx/txt)."""
    pieces = []
    file_sources = []
    if not uploaded_files:
        return "", []
    for f in uploaded_files:
        name = getattr(f, "name", "uploaded_file")
        mime = getattr(f, "type", "")
        try:
            if name.lower().endswith(".pdf") or mime == "application/pdf":
                reader = PyPDF2.PdfReader(f)
                text = []
                for p in reader.pages:
                    page_text = p.extract_text() or ""
                    text.append(page_text)
                file_text = "\n".join(text)
                pieces.append(f"[File: {name}]\n{file_text}\n")
                file_sources.append(f"{name}")
            elif name.lower().endswith(".docx") or mime in (
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            ):
                doc = Document(f)
                text = []
                for para in doc.paragraphs:
                    if para.text:
                        text.append(para.text)
                file_text = "\n".join(text)
                pieces.append(f"[File: {name}]\n{file_text}\n")
                file_sources.append(f"{name}")
            elif name.lower().endswith(".txt") or mime == "text/plain":
                raw = f.read()
                if isinstance(raw, bytes):
                    raw = raw.decode("utf-8", errors="ignore")
                pieces.append(f"[File: {name}]\n{raw}\n")
                file_sources.append(f"{name}")
            else:
                # fallback: try to read as text
                raw = f.read()
                if isinstance(raw, bytes):
                    raw = raw.decode("utf-8", errors="ignore")
                pieces.append(f"[File: {name}]\n{raw}\n")
                file_sources.append(f"{name}")
        except Exception as e:
            pieces.append(f"[File: {name}] (error reading file: {e})\n")
            file_sources.append(f"{name} (read error)")
    combined = "\n\n".join(pieces)
    return combined, file_sources

# -------------------------
# Mock Pathway Live Ingestion
# -------------------------
def ingest_live_update(title: str, source: str, content: str):
    """Simulate ingesting a live article/blog update into Pathway (local mock)."""
    entry = {
        "id": str(int(time.time()*1000)),
        "title": title,
        "source": source,
        "content": content,
        "ts": time.strftime("%Y-%m-%d %H:%M:%S"),
    }
    st.session_state.live_feed.insert(0, entry)  # newest first



# -------------------------
# LLM -> report generation
# -------------------------
def call_llm_generate(question: str, file_text: str, live_text: str):
    """
    Call the LLM to generate a structured report.
    Returns full text (string).
    If client not configured, returns a mock report.
    """
    system_prompt = (
        "You are an expert research assistant. Generate a structured, evidence-based research report "
        "that contains: Key Takeaways (bulleted), Abstract, Introduction, Main Sections depending on the question, "
        "Conclusion, and References. Inline-cite sources using [1], [2] etc. At the end include a 'Sources' section "
        "that maps citation numbers to source names/URLs/pages.\n\n"
        "If provided with uploaded file content or live feed content, use that content as primary evidence. "
        "If content isn't provided, produce a concise general report."
    )

    user_prompt = f"""Question: {question}

Uploaded file content (if any):
{file_text or '[none]'}

Live feed content (if any):
{live_text or '[none]'}

Instructions:
- Compose a report ~ 400-800 words depending on complexity.
- Use inline citation markers like [1], [2] where you reference the provided content.
- At the end, include a "Sources:" section listing sources in the format:
  [1] source description (e.g., 'myfile.pdf p.12' or 'MockNews: article title (2025-09-20)')
"""

    # Mock mode if client unavailable
    if client is None:
        # Create a deterministic mock report that references any sources found
        sources = []
        idx = 1
        if file_text and file_text.strip() != "":
            sources.append(f"[{idx}] Uploaded Files combined (user files)")
            idx += 1
        if live_text and live_text.strip() != "":
            sources.append(f"[{idx}] Live feed updates (ingested)")
            idx += 1
        if not sources:
            sources = ["[1] General knowledge / no sources provided"]

        # Build mock report
        report_lines = []
        report_lines.append(f"# Research Report: {question}\n")
        report_lines.append("## Key Takeaways")
        report_lines.append("- This is a mock key takeaway generated for demo purposes.")
        report_lines.append("- The assistant will use uploaded files and live feed when available.")
        report_lines.append("\n## Abstract")
        report_lines.append("This mock report demonstrates the Smart Research Assistant functionality.")
        report_lines.append("\n## Introduction")
        report_lines.append("The system ingests documents and live feeds, then synthesizes answers with citations.")
        report_lines.append("\n## Detailed Findings")
        report_lines.append("Detailed analysis would come from the LLM in production. Example reference: [1].")
        report_lines.append("\n## Conclusion")
        report_lines.append("Mock conclusion.")
        report_lines.append("\n## Sources:")
        report_lines.extend(sources)
        return "\n\n".join(report_lines)

    # Real LLM call
    try:
        response = client.chat.completions.create(
            model=MODEL_NAME,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ],
            temperature=0.1,
            max_tokens=1500
        )
        # The returned structure depends on the client lib; adapt if necessary
        report_text = response.choices[0].message.content
        return report_text
    except Exception as e:
        # fallback message
        return f"(LLM error: {e})\n\n" + call_llm_generate(question, file_text, live_text)  # try mock

# -------------------------
# UI: Custom CSS (keeps your original theme)
# -------------------------
st.markdown("""
<style>
body, .stApp {
    background: linear-gradient(135deg, #0f2027, #203a43, #2c5364);
    color: #FFFFFF;
    font-family: "Poppins", sans-serif;
}
.header { display: flex; align-items: center; justify-content: space-between; margin-bottom: 10px; }
.header-left { display:flex; align-items:center; gap:12px; }
.logo { font-size: 48px; }
.title { font-size: 30px; font-weight: 800; color: #7dd3fc; }
.tagline { color: #cfe8f6; font-size: 14px; margin-top:2px; }
.stats-box { background-color: rgba(255,255,255,0.05); padding: 10px; border-radius: 10px; min-width: 160px; text-align:right; }
.upload-box { border-radius: 12px; padding: 12px; background-color: rgba(255,255,255,0.03); }
.report-box { background-color: rgba(255,255,255,0.03); padding: 18px; border-radius: 10px; }
</style>
""", unsafe_allow_html=True)

# -------------------------
# Header area
# -------------------------
col1, col2 = st.columns([3, 1])
with col1:
    st.markdown("<div class='header-left'><div class='logo'>🧠</div><div><div class='title'>Smart Research Assistant</div><div class='tagline'>AI-powered research with live data & citations</div></div></div>", unsafe_allow_html=True)
with col2:
    st.markdown(f"<div class='stats-box'>🔋 <b>Credits Used:</b> {st.session_state.credits_used:.2f}<br>🔎 <b>Credits Remaining:</b> {st.session_state.credits_remaining:.2f}</div>", unsafe_allow_html=True)

st.write("---")

# -------------------------
# Input area: upload + question
# -------------------------
with st.container():
    st.markdown("### 📂 Upload your research files (PDF, DOCX, TXT) — optional")
    uploaded_files = st.file_uploader(
        label="Drag & drop files here (200MB max per file)",
        type=["pdf", "docx", "txt"],
        accept_multiple_files=True,
        help="Upload PDFs, Word docs, or text files. Their content will be used to generate the report."
    )

    # show list of uploaded file names
    if uploaded_files:
        names = [f.name for f in uploaded_files]
        st.markdown("**Uploaded:** " + ", ".join(names))
        # small 'clear' button
        if st.button("Remove uploaded files"):
            # Clearing file_uploader client-side requires rerun; reset by setting uploaded_files to None and rerun
            st.experimental_rerun()

    st.markdown("### 💡 Ask your research question")
    question = st.text_input("Enter your question (e.g., 'components of DBMS')", key="question_input")
    gen_col1, gen_col2 = st.columns([1,1])
    with gen_col1:
        generate_clicked = st.button("⚡ Generate Report", key="generate_btn")
    with gen_col2:
        st.write("")  # placeholder to keep layout

# -------------------------
# Left sidebar: Sources & Live Feed
# -------------------------
with st.sidebar:
    st.header("📚 Sources & Live Feed")
    if st.session_state.sources:
        st.subheader("Document Sources")
        for s in st.session_state.sources:
            st.markdown(f"- {s}")
    else:
        st.info("⚠ No document sources found yet. Upload files to populate sources.")

    st.markdown("---")
    
    # show live feed entries
    if st.session_state.live_feed:
        st.subheader("Live Feed Updates")
        for entry in st.session_state.live_feed:
            st.markdown(f"**{entry['title']}**")
            st.markdown(f"*{entry['source']} — {entry['ts']}*")
            st.markdown(entry['content'][:200] + ("..." if len(entry['content'])>200 else ""))
            st.markdown("---")
    else:
        st.info("No live updates yet.")

    # ingest new live update (mock UI)
    st.markdown("**Ingest mock live update**")
    new_title = st.text_input("Title for live update", key="live_title")
    new_source = st.text_input("Source (e.g., blog.example.com)", key="live_source")
    new_content = st.text_area("Content", key="live_content", height=120)
    if st.button("Ingest Live Update", key="ingest_btn"):
        if new_title.strip() and new_content.strip():
            ingest_live_update(new_title.strip(), new_source.strip() or "MockSource", new_content.strip())
            st.success("Live update ingested.")
            st.experimental_rerun()
        else:
            st.error("Please provide title and content for the live update.")

    st.markdown("---")
    st.subheader("Billing Log (mock Flexprice)")
    if st.session_state.billing_log:
        for rec in st.session_state.billing_log[-10:][::-1]:
            st.markdown(f"- {rec['ts']}: \"{rec['question']}\" → {rec['cost']} credit(s)")
    else:
        st.markdown("No billing activity yet.")

# -------------------------
# Handle Generate Report click
# -------------------------
if generate_clicked:
    if not question or question.strip() == "":
        st.error("❌ Please enter a research question.")
    else:
        # Extract file content (if files)
        file_text, file_sources = extract_file_content(uploaded_files) if uploaded_files else ("", [])
        # Compose live feed combined text
        live_text = "\n\n".join([f"{e['title']} ({e['source']}):\n{e['content']}" for e in st.session_state.live_feed])

        # prepare sources list for sidebar display and for mapping citations
        combined_sources = []
        for i, name in enumerate(file_sources, start=1):
            combined_sources.append({"id": i, "desc": f"{name} (uploaded file)"})
        offset = len(combined_sources)
        for j, e in enumerate(st.session_state.live_feed, start=1):
            combined_sources.append({"id": offset + j, "desc": f"{e['source']}: {e['title']} ({e['ts']})"})

        # show a spinner with staged messages
        with st.spinner("🔎 Fetching sources and preparing context..."):
            time.sleep(0.6)

        # Call the LLM (or mock)
        with st.spinner("✍️ Summarizing and generating the report (this may take a few seconds)..."):
            report_text = call_llm_generate(question.strip(), file_text, live_text)
            # small delay for UX
            time.sleep(0.5)

        # Post-process: try to extract Sources section if present
        extracted_sources = []
        if "## Sources:" in report_text or "\n## Sources" in report_text or "\nSources:" in report_text:
            # naive split to find sources - keep lines after "Sources"
            lower = report_text.lower()
            idx = lower.find("sources")
            if idx != -1:
                src_part = report_text[idx:]
                lines = src_part.splitlines()
                for ln in lines:
                    ln = ln.strip()
                    if ln.startswith("[") or ln.startswith("-") or ln.startswith("•"):
                        # treat as bullet
                        extracted_sources.append(ln.strip("-• ").strip())
                # fallback: if none found, leave empty
        # fallback: map our combined_sources if no explicit sources found
        if not extracted_sources and combined_sources:
            extracted_sources = [f"[{s['id']}] {s['desc']}" for s in combined_sources]

        # Update session state counters and billing
        st.session_state.questions += 1
        st.session_state.reports += 1
        st.session_state.credits_used += COST_PER_REPORT
        st.session_state.credits_remaining = max(0.0, st.session_state.credits_remaining - COST_PER_REPORT)
        st.session_state.billing_log.append({
            "question": question.strip(),
            "cost": COST_PER_REPORT,
            "ts": time.strftime("%Y-%m-%d %H:%M:%S")
        })

        # Save last report for download
        st.session_state.last_report = {
            "question": question.strip(),
            "report": report_text,
            "sources": extracted_sources,
            "generated_at": time.strftime("%Y-%m-%d %H:%M:%S")
        }

        # Update sidebar document sources (file_sources)
        if file_sources:
            st.session_state.sources = file_sources

        # Display the report in the main area
        st.markdown("### 📑 Final Report")
        # Use an info box for Key Takeaways if we can extract them
        # Try to find "Key Takeaways" section
        key_takeaways = []
        if "Key Takeaways" in report_text:
            try:
                # naive extraction: split between "Key Takeaways" and next "##"
                start = report_text.find("Key Takeaways")
                tail = report_text[start:]
                # look for next section header
                next_sec_idx = tail.find("\n\n")
                # parse paragraph lines
                lines = tail.splitlines()
                for ln in lines[1:8]:
                    if ln.strip().startswith("-") or ln.strip().startswith("•"):
                        key_takeaways.append(ln.strip("-• ").strip())
            except Exception:
                key_takeaways = []

        if key_takeaways:
            st.markdown("<div class='report-box'><b>🔎 Key Takeaways</b></div>", unsafe_allow_html=True)
            for kt in key_takeaways:
                st.markdown(f"- {kt}")

        # Show collapsible report content
        with st.expander("Show full structured report (expand/collapse)", expanded=True):
            st.markdown(f"<div class='report-box'>{report_text.replace('\\n','  \n')}</div>", unsafe_allow_html=True)

        # Show Sources in a neat box below or aside
        st.markdown("### 📚 Sources (extracted)")
        if extracted_sources:
            for s in extracted_sources:
                st.markdown(f"- {s}")
        else:
            st.info("No explicit sources were found in the generated report.")

        # Download buttons
        col_dl1, col_dl2 = st.columns(2)
        with col_dl1:
            docx_buf = save_to_docx(report_text, title=f"Report - {question.strip()}")
            st.download_button("⬇️ Download as DOCX", data=docx_buf, file_name="research_report.docx",
                               mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        with col_dl2:
            pdf_buf = save_to_pdf(report_text, title=f"Report - {question.strip()}")
            st.download_button("⬇️ Download as PDF", data=pdf_buf, file_name="research_report.pdf", mime="application/pdf")

        st.success("✅ Report generated and saved to session. See Sources & Billing in sidebar.")

# -------------------------
# If there is a last report, show a small history card
# -------------------------
if st.session_state.last_report:
    st.write("---")
    st.markdown("### 🕘 Last generated report")
    lr = st.session_state.last_report
    st.markdown(f"**Question:** {lr['question']}  \n**Generated at:** {lr['generated_at']}")
    st.markdown("**Sources:**")
    for s in lr["sources"]:
        st.markdown(f"- {s}")

# -------------------------
# Footer removed per user request
